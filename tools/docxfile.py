from collections.abc import Generator
from typing import Any
import io
import os
import re
import zipfile
import shutil

from dify_plugin import Tool
from dify_plugin.entities.tool import ToolInvokeMessage
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt
from lxml import etree

# ==== 常量定义 ====
NS_CT = "http://schemas.openxmlformats.org/package/2006/content-types"
NS_REL = "http://schemas.openxmlformats.org/package/2006/relationships"
REL_TYPE_NUMBERING = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering"
REL_TYPE_STYLES = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles"
REL_TYPE_THEME = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme"

PARTS_TO_COPY = [
    "word/styles.xml",
    "word/stylesWithEffects.xml",
    "word/theme/theme1.xml",
    "word/fontTable.xml",
    "word/numbering.xml",
    "word/settings.xml",
    "word/_rels/styles.xml.rels",
    "word/_rels/numbering.xml.rels",
]

CONTENT_TYPES = {
    "/word/styles.xml": "application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml",
    "/word/stylesWithEffects.xml": "application/vnd.ms-word.stylesWithEffects+xml",
    "/word/theme/theme1.xml": "application/vnd.openxmlformats-officedocument.theme+xml",
    "/word/fontTable.xml": "application/vnd.openxmlformats-officedocument.wordprocessingml.fontTable+xml",
    "/word/numbering.xml": "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml",
    "/word/settings.xml": "application/vnd.openxmlformats-officedocument.wordprocessingml.settings+xml",
}


# ==== 工具函数 ====
def unzip_docx(docx_path, extract_dir):
    if os.path.exists(extract_dir):
        shutil.rmtree(extract_dir)
    os.makedirs(extract_dir)
    with zipfile.ZipFile(docx_path, 'r') as zf:
        zf.extractall(extract_dir)


def zip_dir(input_dir, output_docx):
    with zipfile.ZipFile(output_docx, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, _, files in os.walk(input_dir):
            for f in files:
                full = os.path.join(root, f)
                arc = os.path.relpath(full, input_dir)
                zipf.write(full, arc)


def ensure_content_type_override(ct_path, partname, content_type):
    tree = etree.parse(ct_path)
    root = tree.getroot()
    found = root.xpath(f'ct:Override[@PartName="{partname}"]', namespaces={"ct": NS_CT})
    if not found:
        override = etree.Element(f"{{{NS_CT}}}Override")
        override.set("PartName", partname)
        override.set("ContentType", content_type)
        root.append(override)
        tree.write(ct_path, xml_declaration=True, encoding="utf-8", standalone="yes")


def read_relationships(doc_rels_path):
    if not os.path.exists(doc_rels_path):
        os.makedirs(os.path.dirname(doc_rels_path), exist_ok=True)
        root = etree.Element(f"{{{NS_REL}}}Relationships")
        tree = etree.ElementTree(root)
        tree.write(doc_rels_path, xml_declaration=True, encoding="utf-8", standalone="yes")
    tree = etree.parse(doc_rels_path)
    return tree


def ensure_relationship(doc_rels_path, rel_type, target):
    tree = read_relationships(doc_rels_path)
    root = tree.getroot()
    nsmap = {"r": NS_REL}
    exists = root.xpath(
        f'r:Relationship[@Type="{rel_type}" and @Target="{target}"]',
        namespaces=nsmap
    )
    if exists:
        tree.write(doc_rels_path, xml_declaration=True, encoding="utf-8", standalone="yes")
        return

    existing_ids = set([el.get("Id") for el in root.findall(f"{{{NS_REL}}}Relationship") if el.get("Id")])
    next_num = 1
    rid = f"rId{next_num}"
    while rid in existing_ids:
        next_num += 1
        rid = f"rId{next_num}"

    rel = etree.Element(f"{{{NS_REL}}}Relationship")
    rel.set("Id", rid)
    rel.set("Type", rel_type)
    rel.set("Target", target)
    root.append(rel)
    tree.write(doc_rels_path, xml_declaration=True, encoding="utf-8", standalone="yes")


def safe_copy(src_root, dst_root, rel_path):
    src = os.path.join(src_root, rel_path)
    dst = os.path.join(dst_root, rel_path)
    if os.path.exists(src):
        os.makedirs(os.path.dirname(dst), exist_ok=True)
        shutil.copy2(src, dst)
        return True
    return False


def copy_styles_and_dependencies_from_template(template_path, output_path):
    # 新建空白docx
    Document().save(output_path)

    # 解压模板和空白文档
    t_dir = "_tmp_tpl_unzip"
    b_dir = "_tmp_blank_unzip"
    unzip_docx(template_path, t_dir)
    unzip_docx(output_path, b_dir)

    # 复制样式相关部件
    for part in PARTS_TO_COPY:
        safe_copy(t_dir, b_dir, part)

    # 修正Content_Types.xml
    ct_path = os.path.join(b_dir, "[Content_Types].xml")
    for partname, ctype in CONTENT_TYPES.items():
        if os.path.exists(os.path.join(b_dir, partname.lstrip("/"))):
            ensure_content_type_override(ct_path, partname, ctype)

    # 确保document.xml.rels关系
    doc_rels_path = os.path.join(b_dir, "word/_rels/document.xml.rels")
    ensure_relationship(doc_rels_path, REL_TYPE_STYLES, "styles.xml")
    ensure_relationship(doc_rels_path, REL_TYPE_THEME, "theme/theme1.xml")
    if os.path.exists(os.path.join(b_dir, "word/numbering.xml")):
        ensure_relationship(doc_rels_path, REL_TYPE_NUMBERING, "numbering.xml")

    # 重新打包
    zip_dir(b_dir, output_path)

    # 清理临时文件
    shutil.rmtree(t_dir, ignore_errors=True)
    shutil.rmtree(b_dir, ignore_errors=True)


# ==== 主工具类 ====
class DocxWithTemplateStyleTool(Tool):
    def _invoke(self, tool_parameters: dict[str, Any]) -> Generator[ToolInvokeMessage]:
        # def generate_docx_with_template(template_bytes, markdown_content):
        #     """从模板字节流和markdown内容生成目标docx字节流"""
        #     # 临时文件路径
        #     tmp_template = "tmp_template.docx"
        #     tmp_output = "tmp_output.docx"
        #
        #     try:
        #         # 保存模板字节流为临时文件
        #         with open(tmp_template, "wb") as f:
        #             f.write(template_bytes)
        #
        #         # 复制模板样式到空白docx
        #         copy_styles_and_dependencies_from_template(tmp_template, tmp_output)
        #
        #         # 加载带样式的docx并解析markdown写入内容
        #         target_doc = Document(tmp_output)
        #         lines = markdown_content.split('\n')
        #         i = 0
        #         while i < len(lines):
        #             print("this is i , i = " + str(i))
        #             line = lines[i].strip()
        #             # 标题处理
        #             if line.startswith('<div class="title"'):
        #                 match = re.search(r'>(.*?)</div>', line)
        #                 if match:
        #                     print("-------------------this is title---------------------------------------")
        #                     p = target_doc.add_paragraph()
        #                     p.style = "TOC 标题2" if "TOC 标题2" in target_doc.styles else target_doc.styles['Title']
        #                     p.add_run(match.group(1))
        #             # 普通文本处理
        #             elif line.startswith('<div class="text"'):
        #                 match = re.search(r'>(.*?)</div>', line)
        #                 if match:
        #                     print("----------------------this is text--------------------------------")
        #                     p = target_doc.add_paragraph()
        #                     p.style = "正文段落文本" if "正文段落文本" in target_doc.styles else target_doc.styles['Normal']
        #                     p.add_run(match.group(1))
        #             # 一级标题处理
        #             elif line.startswith('<div class="onetitle"'):
        #                 match = re.search(r'>(.*?)</div>', line)
        #                 if match:
        #                     p = target_doc.add_paragraph()
        #                     p.style = "标题2" if "标题2" in target_doc.styles else target_doc.styles['Heading 1']
        #                     p.add_run(match.group(1))
        #             # 二级标题处理
        #             elif line.startswith('<div class="twotitle"'):
        #                 match = re.search(r'>(.*?)</div>', line)
        #                 if match:
        #                     p = target_doc.add_paragraph()
        #                     p.style = "样式11111" if "样式11111" in target_doc.styles else target_doc.styles['Heading 2']
        #                     p.add_run(match.group(1))
        #             # 三级标题处理
        #             elif line.startswith('<div class="threetitle"'):
        #                 match = re.search(r'>(.*?)</div>', line)
        #                 if match:
        #                     p = target_doc.add_paragraph()
        #                     p.style = "标题3" if "标题3" in target_doc.styles else target_doc.styles['Heading 3']
        #                     p.add_run(match.group(1))
        #             # 表格处理
        #             elif line.startswith('<table>'):
        #                 table_data = []
        #                 i += 1
        #                 table_content = []
        #                 while i < len(lines) and not lines[i].strip().startswith('</table>'):
        #                     table_content.append(lines[i])
        #                     i += 1
        #                 table_text = '\n'.join(table_content)
        #
        #                 # 提取表头
        #                 header_match = re.search(r'<thead>(.*?)</thead>', table_text, re.DOTALL)
        #                 if header_match:
        #                     header_content = header_match.group(1)
        #                     header_row = re.findall(r'<th>(.*?)</th>', header_content, re.DOTALL)
        #                     if header_row:
        #                         table_data.append([cell.strip() for cell in header_row])
        #
        #                 # 提取表体
        #                 body_match = re.search(r'<tbody>(.*?)</tbody>', table_text, re.DOTALL)
        #                 if body_match:
        #                     body_content = body_match.group(1)
        #                     rows = re.findall(r'<tr>(.*?)</tr>', body_content, re.DOTALL)
        #                     for row in rows:
        #                         cells = re.findall(r'<td>(.*?)</td>', row, re.DOTALL)
        #                         if cells:
        #                             table_data.append([cell.strip() for cell in cells])
        #
        #                 # 创建表格
        #                 if len(table_data) > 1:
        #                     table = target_doc.add_table(rows=1, cols=len(table_data[0]))
        #                     table.style = 'Table Grid'
        #
        #                     # 设置表格字体
        #                     for row in table.rows:
        #                         for cell in row.cells:
        #                             for paragraph in cell.paragraphs:
        #                                 for run in paragraph.runs:
        #                                     run.font.size = Pt(10)
        #
        #                     # 设置表头
        #                     hdr_cells = table.rows[0].cells
        #                     for j, cell in enumerate(table_data[0]):
        #                         hdr_cells[j].text = cell
        #                         for paragraph in hdr_cells[j].paragraphs:
        #                             paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #                             for run in paragraph.runs:
        #                                 run.bold = True
        #                                 run.font.size = Pt(10)
        #                         # 表头背景色
        #                         shading_elm = parse_xml(r'<w:shd {} w:fill="5B9BD5"/>'.format(nsdecls('w')))
        #                         hdr_cells[j]._tc.get_or_add_tcPr().append(shading_elm)
        #
        #                     # 添加数据行
        #                     for row in table_data[1:]:
        #                         row_cells = table.add_row().cells
        #                         for j, cell in enumerate(row):
        #                             row_cells[j].text = cell
        #                             for paragraph in row_cells[j].paragraphs:
        #                                 paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        #                                 for run in paragraph.runs:
        #                                     run.font.size = Pt(10)
        #             # 其他非空行处理
        #             elif line != '' and not line.startswith('</'):
        #                 print("----------------------this is text,too--------------------------------")
        #                 p = target_doc.add_paragraph()
        #                 p.style = "正文段落文本" if "正文段落文本" in target_doc.styles else target_doc.styles['Normal']
        #                 p.add_run(line)
        #
        #             i += 1
        #
        #         # 将docx保存到字节流
        #         doc_bytes = io.BytesIO()
        #         target_doc.save(doc_bytes)
        #         doc_bytes.seek(0)
        #         return doc_bytes.getvalue()
        #
        #     except Exception as e:
        #         print(f"处理文档时发生未知错误: {str(e)}")
        #
        #     finally:
        #         # 清理临时文件
        #         if os.path.exists(tmp_template):
        #             os.remove(tmp_template)
        #         if os.path.exists(tmp_output):
        #             os.remove(tmp_output)

        def generate_docx_with_template(template_bytes, markdown_content):
            """从模板字节流和markdown内容生成目标docx字节流"""
            # 临时文件路径
            tmp_template = "tmp_template.docx"
            tmp_output = "tmp_output.docx"

            try:
                # 保存模板字节流为临时文件
                with open(tmp_template, "wb") as f:
                    f.write(template_bytes)

                # 复制模板样式到空白docx
                copy_styles_and_dependencies_from_template(tmp_template, tmp_output)

                # 加载带样式的docx并解析markdown写入内容
                target_doc = Document(tmp_output)
                lines = markdown_content.split('\n')
                i = 0
                while i < len(lines):
                    print("this is i , i = " + str(i))
                    line = lines[i].strip()
                    # 标题处理
                    if line.startswith('<div class="title"'):
                        match = re.search(r'>(.*?)</div>', line)
                        if match:
                            p = target_doc.add_paragraph()
                            p.style = "A标题"
                            p.add_run(match.group(1))
                    # 普通文本处理
                    elif line.startswith('<div class="text"'):
                        match = re.search(r'>(.*?)</div>', line)
                        if match:
                            p = target_doc.add_paragraph()
                            p.style = "A正文"
                            p.add_run(match.group(1))
                    # 一级标题处理
                    elif line.startswith('<div class="onetitle"'):
                        match = re.search(r'>(.*?)</div>', line)
                        if match:
                            p = target_doc.add_paragraph()
                            p.style = "A一级标题"
                            p.add_run(match.group(1))
                    # 二级标题处理
                    elif line.startswith('<div class="twotitle"'):
                        match = re.search(r'>(.*?)</div>', line)
                        if match:
                            p = target_doc.add_paragraph()
                            p.style = "A二级标题"
                            p.add_run(match.group(1))
                    # 三级标题处理
                    elif line.startswith('<div class="threetitle"'):
                        match = re.search(r'>(.*?)</div>', line)
                        if match:
                            p = target_doc.add_paragraph()
                            p.style = "A三级标题"
                            p.add_run(match.group(1))
                    # 表格处理 - 修复版本
                    elif line.startswith('<table>'):
                        table_data = []
                        i += 1  # 跳过 <table> 行
                        table_content = []

                        # 收集表格内容，直到遇到 </table>
                        while i < len(lines) and not lines[i].strip().startswith('</table>'):
                            table_content.append(lines[i])
                            i += 1

                        # 现在 i 指向 </table> 行，我们需要处理这个表格
                        table_text = '\n'.join(table_content)
                        print(f"表格内容长度: {len(table_text)}")

                        # 提取表头
                        header_match = re.search(r'<thead>(.*?)</thead>', table_text, re.DOTALL)
                        if header_match:
                            header_content = header_match.group(1)
                            header_row = re.findall(r'<th>(.*?)</th>', header_content, re.DOTALL)
                            if header_row:
                                table_data.append([cell.strip() for cell in header_row])
                                print(f"提取到表头: {header_row}")

                        # 提取表体
                        body_match = re.search(r'<tbody>(.*?)</tbody>', table_text, re.DOTALL)
                        if body_match:
                            body_content = body_match.group(1)
                            rows = re.findall(r'<tr>(.*?)</tr>', body_content, re.DOTALL)
                            for row in rows:
                                # 同时匹配 th 和 td，因为您的表格中表体也使用了 th
                                cells = re.findall(r'<(?:th|td)>(.*?)</(?:th|td)>', row, re.DOTALL)
                                if cells:
                                    table_data.append([cell.strip() for cell in cells])
                                    print(f"提取到行数据: {cells}")

                        # 创建表格
                        if table_data and len(table_data) > 0:
                            print(f"创建表格，行数: {len(table_data)}，列数: {len(table_data[0])}")
                            table = target_doc.add_table(rows=1, cols=len(table_data[0]))
                            table.style = 'Table Grid'

                            # 设置表格字体
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.size = Pt(10)

                            # 设置表头
                            hdr_cells = table.rows[0].cells
                            for j, cell in enumerate(table_data[0]):
                                hdr_cells[j].text = cell
                                for paragraph in hdr_cells[j].paragraphs:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for run in paragraph.runs:
                                        run.bold = True
                                        run.font.size = Pt(10)
                                # 表头背景色
                                shading_elm = parse_xml(r'<w:shd {} w:fill="5B9BD5"/>'.format(nsdecls('w')))
                                hdr_cells[j]._tc.get_or_add_tcPr().append(shading_elm)

                            # 添加数据行
                            for row in table_data[1:]:
                                row_cells = table.add_row().cells
                                for j, cell in enumerate(row):
                                    row_cells[j].text = cell
                                    for paragraph in row_cells[j].paragraphs:
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        for run in paragraph.runs:
                                            run.font.size = Pt(10)
                    # 其他非空行处理
                    elif line != '' and not line.startswith('</'):
                        p = target_doc.add_paragraph()
                        p.style = "A正文"
                        p.add_run(line)

                    i += 1

                # 将docx保存到字节流
                doc_bytes = io.BytesIO()
                target_doc.save(doc_bytes)
                doc_bytes.seek(0)
                return doc_bytes.getvalue()

            except Exception as e:
                print(f"处理文档时发生未知错误: {str(e)}")
                import traceback
                traceback.print_exc()

            finally:
                # 清理临时文件
                if os.path.exists(tmp_template):
                    os.remove(tmp_template)
                if os.path.exists(tmp_output):
                    os.remove(tmp_output)



        # 获取输入参数
        template_file = tool_parameters['model']  # Dify的文件对象
        markdown_content = tool_parameters['query']
        output_filename = tool_parameters.get('output_filename', '高校毕业生就业情况分析报告.docx')
        if not output_filename.lower().endswith('.docx'):
            output_filename += '.docx'

        try:
            # 关键修改：从Dify文件对象获取字节流（通过blob属性）
            if hasattr(template_file, 'blob'):
                # 优先使用blob属性获取字节流（符合Dify文件处理规范）
                template_bytes = template_file.blob
            elif hasattr(template_file, 'read'):
                # 兼容其他文件对象类型
                template_bytes = template_file.read()
            elif isinstance(template_file, bytes):
                # 直接传入字节流的情况
                template_bytes = template_file
            else:
                raise TypeError("模板文件格式错误，需要Dify文件对象或字节流")

            # 验证字节流有效性（尝试解压验证是否为docx）
            try:
                with zipfile.ZipFile(io.BytesIO(template_bytes)) as test_zip:
                    # 检查docx必需的核心文件
                    required_files = ['[Content_Types].xml', 'word/document.xml']
                    if not all(f in test_zip.namelist() for f in required_files):
                        raise ValueError("模板文件不是有效的DOCX格式")
            except zipfile.BadZipFile:
                raise ValueError("模板文件不是有效的ZIP压缩文件（可能不是DOCX）")

            # 生成最终docx
            final_docx_bytes = generate_docx_with_template(template_bytes, markdown_content)
            # final_docx_bytes = template_file.blob

            # 输出结果
            yield self.create_blob_message(
                blob=final_docx_bytes,
                meta={
                    'mime_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'filename': output_filename
                }
            )

        except Exception as e:
            yield self.create_text_message(f"生成带模板样式的docx时出错: {str(e)}")